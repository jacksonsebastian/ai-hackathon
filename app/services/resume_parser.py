"""
Resume parsing service.

Handles PDF/DOCX upload, text extraction, and LLM-powered
structured data extraction into candidate profiles.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

from app.database.models import Resume
from app.database import crud
from app.prompts.templates import build_resume_parse_prompt
from app.services.ai_service import AIService
from app.utils.helpers import generate_id, now_utc, hash_content, extract_json_from_response
from app.utils.logger import get_service_logger

logger = get_service_logger("resume_parser")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("PyMuPDF not installed. Install with: pip install pymupdf")
        return "[PDF extraction requires PyMuPDF. Install with: pip install pymupdf]"
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"[Error extracting PDF: {e}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        return "[DOCX extraction requires python-docx. Install with: pip install python-docx]"
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return f"[Error extracting DOCX: {e}]"


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a resume file based on extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF or DOCX.")


async def parse_resume(
    file_bytes: bytes,
    filename: str,
    ai_service: Optional[AIService] = None,
) -> Resume:
    """
    Full resume parsing pipeline:
    1. Extract raw text from PDF/DOCX
    2. Use LLM to extract structured data (primary)
    3. Fall back to regex if LLM fails
    4. Store in database
    5. Return Resume model
    """
    import re
    logger.info(f"Parsing resume: {filename}")

    # Step 1: Extract raw text
    raw_text = extract_text(file_bytes, filename)
    if not raw_text or raw_text.startswith("["):
        logger.warning(f"Limited text extraction for {filename}")

    logger.info(f"Extracted {len(raw_text)} characters from {filename}")

    # Step 2: Check for duplicates
    content_hash = hash_content(raw_text)

    # Step 3: LLM-powered structured extraction (PRIMARY)
    result = {}
    try:
        service = ai_service or AIService()
        prompt = build_resume_parse_prompt(raw_text)
        
        result = await service.generate_structured(
            prompt=prompt,
            system_prompt="You are an expert resume parser. Return ONLY valid JSON. No explanations, no markdown, no code blocks.",
        )
        logger.info(f"LLM parse result keys: {list(result.keys())}")
    except Exception as e:
        logger.error(f"LLM resume parsing failed: {e}")
        result = {}

    # Step 4: If LLM failed, enrich with regex as safety net
    if (
        "raw_response" in result
        or result.get("candidate_name", "Unknown") == "Unknown"
        or not result.get("candidate_name")
    ):
        logger.info("LLM parsing incomplete — enriching with regex")
        
        if not result.get("candidate_name") or result.get("candidate_name") == "Unknown":
            for line in raw_text.split('\n'):
                line = line.strip()
                if line and 2 < len(line) < 60:
                    if '@' not in line and 'http' not in line and not re.match(r'^[\d\+\(]', line):
                        if line.upper() not in {"SKILLS", "EXPERIENCE", "EDUCATION", "PROJECTS", "SUMMARY", "OBJECTIVE", "CERTIFICATIONS"}:
                            result["candidate_name"] = line
                            break
        
        if not result.get("email"):
            em = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', raw_text)
            if em:
                result["email"] = em.group(0)
        
        if not result.get("phone"):
            ph = re.search(r'[\+]?[(]?\d{1,4}[)]?[-\s./\d]{7,15}', raw_text)
            if ph:
                result["phone"] = ph.group(0).strip()
        
        if not result.get("skills"):
            skills_list = [
                "Python", "JavaScript", "TypeScript", "Java", "C++", "React",
                "Next.js", "Node.js", "Express", "Django", "Flask", "FastAPI",
                "Docker", "Kubernetes", "AWS", "SQL", "PostgreSQL", "MongoDB",
                "Redis", "Git", "HTML", "CSS", "REST", "GraphQL", "Tailwind",
            ]
            tl = raw_text.lower()
            result["skills"] = [s for s in skills_list if s.lower() in tl]

    # Step 5: Build Resume model
    ext = Path(filename).suffix.lower().lstrip(".")
    resume = Resume(
        id=generate_id(),
        filename=filename,
        file_type=ext,
        raw_text=raw_text,
        candidate_name=result.get("candidate_name", "Unknown"),
        email=result.get("email", ""),
        phone=result.get("phone", ""),
        skills=result.get("skills", []),
        experience=result.get("experience", []),
        education=result.get("education", []),
        projects=result.get("projects", []),
        certifications=result.get("certifications", []),
        technologies=result.get("technologies", []),
        summary=result.get("summary", raw_text[:500]),
        strengths=result.get("strengths", []),
        gaps=result.get("gaps", []),
        content_hash=content_hash,
        created_at=now_utc(),
        updated_at=now_utc(),
    )

    # Step 6: Store in database
    saved_resume = crud.create_resume(resume)
    logger.info(f"Resume parsed: {saved_resume.candidate_name} (id={saved_resume.id})")

    return saved_resume



def parse_resume_basic(file_bytes: bytes, filename: str) -> Resume:
    """
    Basic synchronous resume parsing without LLM.
    Uses regex to extract name, email, phone, and skills from raw text.
    """
    import re
    
    raw_text = extract_text(file_bytes, filename)
    ext = Path(filename).suffix.lower().lstrip(".")
    
    # Extract email
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', raw_text)
    email = email_match.group(0) if email_match else ""
    
    # Extract phone
    phone_match = re.search(r'[\+]?[(]?[0-9]{1,4}[)]?[-\s./0-9]{7,15}', raw_text)
    phone = phone_match.group(0).strip() if phone_match else ""
    
    # Extract candidate name (first non-empty line that's not an email/phone/url)
    candidate_name = "Unknown"
    for line in raw_text.split('\n'):
        line = line.strip()
        if line and len(line) > 2 and len(line) < 60:
            # Skip lines that look like emails, phones, URLs
            if '@' not in line and 'http' not in line and not re.match(r'^[\d\+\(]', line):
                candidate_name = line
                break
    
    # Extract skills using common tech keywords
    common_skills = [
        "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust", "Ruby",
        "React", "Angular", "Vue", "Node.js", "Express", "Django", "Flask", "FastAPI",
        "Spring", "Docker", "Kubernetes", "AWS", "Azure", "GCP", "SQL", "PostgreSQL",
        "MongoDB", "Redis", "Git", "Linux", "REST", "GraphQL", "HTML", "CSS",
        "TensorFlow", "PyTorch", "Machine Learning", "Deep Learning", "AI",
        "Agile", "Scrum", "CI/CD", "DevOps", "Microservices", "System Design",
    ]
    found_skills = []
    text_lower = raw_text.lower()
    for skill in common_skills:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    
    resume = Resume(
        id=generate_id(),
        filename=filename,
        file_type=ext,
        raw_text=raw_text,
        candidate_name=candidate_name,
        email=email,
        phone=phone,
        skills=found_skills,
        summary=f"Resume parsed from {filename}. {len(found_skills)} skills detected.",
        content_hash=hash_content(raw_text),
        created_at=now_utc(),
        updated_at=now_utc(),
    )
    
    saved_resume = crud.create_resume(resume)
    return saved_resume

